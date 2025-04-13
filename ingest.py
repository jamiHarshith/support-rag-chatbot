import json
from llama_index.core import Document
from pathlib import Path
from docx import Document as DocxDocument
import fitz

def load_web_docs(json_path="data/support_docs.json"):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return [Document(text=item["content"], metadata={"url": item["url"]}) for item in data]


def load_pdfs(pdf_dir="data/InsurancePDFs"):
    docs = []
    for path in Path(pdf_dir).glob("*.pdf"):
        with fitz.open(path) as pdf:
            text = "\n".join([page.get_text() for page in pdf])
            if text.strip():
                docs.append(Document(text=text, metadata={"file": str(path)}))
    return docs

def load_docs(doc_dir="data/InsurancePDFs"):
    docs = []
    for path in Path(doc_dir).glob("*.docx"):
        doc = DocxDocument(path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if text:
            docs.append(Document(text=text, metadata={"file": str(path)}))
    return docs

def load_all():
    web = load_web_docs()
    pdfs = load_pdfs()
    docs = load_docs()
    print(f"Loaded {len(web)} web docs, {len(pdfs)} PDFs, {len(docs)} Word docs")
    return web + pdfs + docs
